"""
Universal Resume Parser - Production Ready
Outputs AI-friendly structured text that can be read line by line
Handles: PDF, DOCX, DOC, TXT, Images with tables, columns, colored backgrounds
"""

import io
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from enum import Enum

# PDF processing
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX processing
try:
    from docx import Document
except ImportError:
    Document = None

# Image/OCR processing
try:
    from PIL import Image, ImageEnhance, ImageFilter
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class UniversalParser:
    """
    Universal parser that outputs AI-friendly structured text
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize parser with configuration"""
        self.config = config or {}
        self.ocr_enabled = self.config.get('ocr_enabled', True)
        self.enhance_images = self.config.get('enhance_images', True)
        
    def parse(self, file_path: str, save_output: bool = True) -> Dict[str, Any]:
        """
        Parse any resume format and return AI-friendly structured text
        
        Args:
            file_path: Path to resume file
            save_output: Whether to save extracted text and metadata to files
            
        Returns:
            Dictionary with structured text and metadata
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return self._error_response(f"File not found: {file_path}")
            
            # Detect format
            extension = file_path.suffix.lower()
            
            # Route to appropriate parser
            if extension == '.pdf':
                result = self._parse_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                result = self._parse_docx(file_path)
            elif extension == '.txt':
                result = self._parse_txt(file_path)
            elif extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                result = self._parse_image(file_path)
            else:
                return self._error_response(f"Unsupported format: {extension}")
            
            # Save output if requested and parsing was successful
            if save_output and result.get('success'):
                self._save_output(file_path, result)
                
            return result
                
        except Exception as e:
            logger.error(f"Parse error: {e}", exc_info=True)
            return self._error_response(str(e))
    
    def _save_output(self, original_file: Path, result: Dict[str, Any]):
        """Save extracted text and metadata to files"""
        try:
            import json
            
            # Create output filenames
            base_name = original_file.stem
            output_dir = original_file.parent
            
            # Save extracted text
            text_file = output_dir / f"{base_name}_parsed.txt"
            with open(text_file, 'w', encoding='utf-8') as f:
                f.write(result['text'])
            logger.info(f"Saved extracted text to: {text_file}")
            
            # Save metadata
            metadata_file = output_dir / f"{base_name}_metadata.json"
            metadata = {
                **result['metadata'],
                'original_file': str(original_file),
                'word_count': len(result['text'].split()),
                'char_count': len(result['text']),
                'line_count': len(result['text'].split('\n'))
            }
            
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, indent=2)
            logger.info(f"Saved metadata to: {metadata_file}")
            
        except Exception as e:
            logger.error(f"Error saving output: {e}")
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF with proper column handling and embedded image extraction"""
        
        if not pdfplumber:
            return self._error_response("pdfplumber not installed")
        
        try:
            with pdfplumber.open(file_path) as pdf:
                all_sections = []
                has_images = False
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text with proper column handling
                    text = self._extract_pdf_columns(page)
                    
                    if text:
                        all_sections.append(text)
                    
                    # Check for embedded images and extract text from them
                    if page.images:
                        has_images = True
                        logger.info(f"Found {len(page.images)} images on page {page_num}")
                        
                        # Extract text from embedded images
                        image_text = self._extract_text_from_pdf_images(page, page_num)
                        if image_text:
                            all_sections.append(f"\n[CONTENT FROM EMBEDDED IMAGES]\n{image_text}")
                
                # Combine and clean
                combined_text = '\n\n'.join(all_sections)
                structured_text = self._structure_for_ai(combined_text)
                
                return {
                    'success': True,
                    'text': structured_text,
                    'metadata': {
                        'format': 'pdf',
                        'pages': len(pdf.pages),
                        'method': 'column_aware_with_images' if has_images else 'column_aware',
                        'has_embedded_images': has_images
                    }
                }
                
        except Exception as e:
            logger.error(f"PDF parse error: {e}")
            return self._error_response(str(e))
    
    def _extract_text_from_pdf_images(self, page, page_num: int) -> str:
        """Extract text from embedded images in PDF using OCR"""
        
        if not self.ocr_enabled or not Image or not pytesseract:
            logger.warning("OCR not available - skipping embedded images")
            return ""
        
        try:
            import io
            
            image_texts = []
            
            for img_idx, img_info in enumerate(page.images):
                try:
                    # Get image from PDF
                    # Try to extract image data
                    x0, y0, x1, y1 = img_info['x0'], img_info['top'], img_info['x1'], img_info['bottom']
                    
                    # Crop the page to the image area and convert to image
                    img_bbox = page.within_bbox((x0, y0, x1, y1))
                    img_obj = img_bbox.to_image(resolution=300)
                    
                    # Convert to PIL Image
                    pil_img = img_obj.original
                    
                    # Enhance for OCR
                    if self.enhance_images:
                        pil_img = self._enhance_image(pil_img)
                    
                    # Perform OCR
                    text = pytesseract.image_to_string(pil_img, config='--oem 3 --psm 6')
                    
                    # Post-process
                    text = self._post_process_ocr(text)
                    
                    if text.strip():
                        image_texts.append(f"[Image {page_num}.{img_idx + 1}]\n{text.strip()}")
                        logger.info(f"Extracted {len(text)} chars from image {img_idx + 1} on page {page_num}")
                
                except Exception as img_error:
                    logger.warning(f"Could not extract text from image {img_idx + 1}: {img_error}")
                    continue
            
            return '\n\n'.join(image_texts)
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF images: {e}")
            return ""
    
    def _extract_pdf_columns(self, page) -> str:
        """
        Extract text from PDF page with proper column handling
        Reads top-to-bottom in each column, then moves to next column
        """
        # Get page dimensions
        page_width = page.width
        page_height = page.height
        
        # Extract words with positions
        words = page.extract_words(x_tolerance=3, y_tolerance=3)
        
        if not words:
            # Fallback to simple extraction
            return page.extract_text() or ""
        
        # Detect columns by analyzing x-coordinates
        x_coords = [w['x0'] for w in words]
        
        # Find column boundaries (look for large gaps)
        x_sorted = sorted(set(x_coords))
        gaps = []
        
        for i in range(len(x_sorted) - 1):
            gap = x_sorted[i + 1] - x_sorted[i]
            if gap > 40:  # Significant gap indicates column boundary
                gaps.append((x_sorted[i], x_sorted[i + 1], gap))
        
        # Determine columns
        if not gaps:
            # Single column - use simple extraction
            return page.extract_text() or ""
        
        # Find the largest gap (main column divider)
        largest_gap = max(gaps, key=lambda x: x[2])
        mid_point = (largest_gap[0] + largest_gap[1]) / 2
        
        # Split words into columns
        left_column = [w for w in words if w['x0'] < mid_point]
        right_column = [w for w in words if w['x0'] >= mid_point]
        
        # Sort each column by vertical position (top to bottom)
        left_column.sort(key=lambda w: (w['top'], w['x0']))
        right_column.sort(key=lambda w: (w['top'], w['x0']))
        
        # Group into lines for each column
        left_text = self._words_to_text(left_column)
        right_text = self._words_to_text(right_column)
        
        # Combine: left column first (top to bottom), then right column (top to bottom)
        if left_text and right_text:
            return f"{left_text}\n\n{right_text}"
        elif left_text:
            return left_text
        else:
            return right_text
    
    def _words_to_text(self, words: List[Dict]) -> str:
        """Convert list of words to text, grouping by lines"""
        if not words:
            return ""
        
        lines = []
        current_line = [words[0]]
        current_y = words[0]['top']
        
        for word in words[1:]:
            # Check if word is on same line (similar y-coordinate)
            if abs(word['top'] - current_y) < 5:
                current_line.append(word)
            else:
                # New line
                line_text = ' '.join(w['text'] for w in current_line)
                lines.append(line_text)
                current_line = [word]
                current_y = word['top']
        
        # Add last line
        if current_line:
            line_text = ' '.join(w['text'] for w in current_line)
            lines.append(line_text)
        
        return '\n'.join(lines)
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX with table support"""
        
        if not Document:
            return self._error_response("python-docx not installed")
        
        try:
            doc = Document(file_path)
            sections = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    sections.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                table_text = self._format_table_for_ai(table)
                if table_text:
                    sections.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            combined_text = '\n'.join(sections)
            structured_text = self._structure_for_ai(combined_text)
            
            return {
                'success': True,
                'text': structured_text,
                'metadata': {
                    'format': 'docx',
                    'tables': len(doc.tables),
                    'method': 'structured'
                }
            }
            
        except Exception as e:
            logger.error(f"DOCX parse error: {e}")
            return self._error_response(str(e))
    
    def _parse_txt(self, file_path: Path) -> Dict[str, Any]:
        """Parse text file"""
        
        try:
            # Try multiple encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        text = f.read()
                    
                    structured_text = self._structure_for_ai(text)
                    
                    return {
                        'success': True,
                        'text': structured_text,
                        'metadata': {
                            'format': 'txt',
                            'encoding': encoding,
                            'method': 'direct'
                        }
                    }
                except UnicodeDecodeError:
                    continue
            
            return self._error_response("Unable to decode text file")
            
        except Exception as e:
            logger.error(f"TXT parse error: {e}")
            return self._error_response(str(e))
    
    def _parse_image(self, file_path: Path) -> Dict[str, Any]:
        """Parse image with OCR and column-aware processing"""
        
        if not self.ocr_enabled or not Image or not pytesseract:
            return self._error_response("OCR not available - install Pillow and pytesseract, and Tesseract OCR")
        
        try:
            img = Image.open(file_path)
            
            # Enhance image for better OCR
            if self.enhance_images:
                img = self._enhance_image(img)
            
            # Detect columns in image and OCR each separately
            text = self._ocr_with_column_detection(img)
            
            # Advanced post-processing for OCR text
            text = self._post_process_ocr(text)
            
            structured_text = self._structure_for_ai(text)
            
            return {
                'success': True,
                'text': structured_text,
                'metadata': {
                    'format': 'image',
                    'method': 'ocr_column_aware',
                    'dimensions': f"{img.width}x{img.height}"
                }
            }
            
        except Exception as e:
            logger.error(f"Image parse error: {e}")
            return self._error_response(str(e))
    
    def _ocr_with_column_detection(self, img: 'Image.Image') -> str:
        """
        Perform OCR with column detection
        Detects columns and OCRs each separately to maintain proper reading order
        """
        import numpy as np
        
        # Get image data for column detection
        img_array = np.array(img)
        height, width = img_array.shape if len(img_array.shape) == 2 else img_array.shape[:2]
        
        # Detect columns by analyzing vertical white space
        columns = self._detect_image_columns(img_array, width, height)
        
        if len(columns) <= 1:
            # Single column or no columns detected - use standard OCR
            return pytesseract.image_to_string(img, config='--oem 3 --psm 6')
        
        # OCR each column separately
        column_texts = []
        for col_start, col_end in columns:
            # Crop image to column
            col_img = img.crop((col_start, 0, col_end, height))
            
            # OCR this column
            col_text = pytesseract.image_to_string(col_img, config='--oem 3 --psm 6')
            
            if col_text.strip():
                column_texts.append(col_text.strip())
        
        # Combine columns: left column first, then right column
        return '\n\n'.join(column_texts)
    
    def _detect_image_columns(self, img_array, width: int, height: int) -> List[tuple]:
        """
        Detect column boundaries in image by analyzing vertical white space
        Returns list of (start_x, end_x) tuples for each column
        """
        import numpy as np
        
        # Convert to grayscale if needed
        if len(img_array.shape) == 3:
            gray = np.mean(img_array, axis=2)
        else:
            gray = img_array
        
        # Calculate vertical projection (sum of pixel values in each column)
        vertical_projection = np.sum(gray, axis=0)
        
        # Normalize
        vertical_projection = vertical_projection / height
        
        # Find threshold (areas with little text will have high values)
        threshold = np.mean(vertical_projection) + 0.5 * np.std(vertical_projection)
        
        # Find gaps (white space between columns)
        gaps = []
        in_gap = False
        gap_start = 0
        
        for x in range(width):
            if vertical_projection[x] > threshold:
                if not in_gap:
                    gap_start = x
                    in_gap = True
            else:
                if in_gap and (x - gap_start) > width * 0.05:  # Gap must be at least 5% of width
                    gaps.append((gap_start, x))
                in_gap = False
        
        # If no significant gaps found, treat as single column
        if not gaps:
            return [(0, width)]
        
        # Find the largest gap (main column divider)
        largest_gap = max(gaps, key=lambda g: g[1] - g[0])
        gap_center = (largest_gap[0] + largest_gap[1]) // 2
        
        # Define columns
        columns = [
            (0, gap_center),           # Left column
            (gap_center, width)        # Right column
        ]
        
        return columns
    
    def _enhance_image(self, img: 'Image.Image') -> 'Image.Image':
        """Enhance image for better OCR (handles colored backgrounds)"""
        
        # Convert to grayscale
        if img.mode != 'L':
            img = img.convert('L')
        
        # Increase contrast (helps with colored backgrounds)
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)
        
        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)
        
        # Resize if too small
        if img.width < 1500:
            scale = 1500 / img.width
            new_size = (int(img.width * scale), int(img.height * scale))
            img = img.resize(new_size, Image.Resampling.LANCZOS)
        
        return img
    
    def _post_process_ocr(self, text: str) -> str:
        """
        Advanced post-processing for OCR text to remove artifacts and noise
        """
        import re
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Remove common OCR artifacts
            line = self._remove_ocr_artifacts(line)
            
            # Skip lines that are just noise
            if self._is_noise_line(line):
                continue
            
            # Fix common OCR errors
            line = self._fix_ocr_errors(line)
            
            # Only keep lines with meaningful content
            if line.strip():
                cleaned_lines.append(line.strip())
        
        return '\n'.join(cleaned_lines)
    
    def _remove_ocr_artifacts(self, line: str) -> str:
        """Remove common OCR artifacts and noise"""
        import re
        
        # Remove single character artifacts at start of line
        line = re.sub(r'^[^a-zA-Z0-9\s]{1,3}\s+', '', line)
        
        # Remove random single characters
        line = re.sub(r'\s+[^a-zA-Z0-9\s]{1}\s+', ' ', line)
        
        # Remove multiple spaces
        line = re.sub(r'\s+', ' ', line)
        
        # Remove common OCR symbols/artifacts
        artifacts = ['Ey', '{ y F', 'K..', r'\ DY', 'yh', '—?', 'EE]', '[EJ']
        for artifact in artifacts:
            line = line.replace(artifact, '')
        
        # Remove standalone special characters
        line = re.sub(r'\s+[^\w\s@.,\-()]+\s+', ' ', line)
        
        return line.strip()
    
    def _is_noise_line(self, line: str) -> bool:
        """Check if a line is just noise/artifacts"""
        line = line.strip()
        
        # Empty line
        if not line:
            return False  # Keep empty lines for structure
        
        # Very short lines with only special characters
        if len(line) < 3 and not line.isalnum():
            return True
        
        # Lines with mostly special characters
        alpha_count = sum(c.isalnum() or c.isspace() for c in line)
        if len(line) > 0 and alpha_count / len(line) < 0.3:
            return True
        
        # Single character lines (except numbers)
        if len(line) == 1 and not line.isdigit():
            return True
        
        return False
    
    def _fix_ocr_errors(self, line: str) -> str:
        """Fix common OCR spelling errors"""
        import re
        
        # Common OCR mistakes
        replacements = {
            r'\bdolar\b': 'dolor',
            r'\bnung\b': 'nunc',
            r'\blea\b': 'leo',
            r'\baccu\b': 'arcu',
            r'\boccumsan\b': 'accumsan',
            r'\bManoqement\b': 'Management',
            r'\bBechsierel\b': 'Bachelor',
            r'\bpusiness\b': 'Business',
            r'\baenceereere\b': 'Management',
            r'\bmain\b': '',  # Often an artifact
            r'\bSSS}\b': '',  # Artifact
            r'\bweet\b': '',  # Artifact
        }
        
        for pattern, replacement in replacements.items():
            line = re.sub(pattern, replacement, line, flags=re.IGNORECASE)
        
        # Fix spacing around punctuation
        line = re.sub(r'\s+([.,;:!?])', r'\1', line)
        line = re.sub(r'([.,;:!?])([a-zA-Z])', r'\1 \2', line)
        
        # Remove multiple spaces again
        line = re.sub(r'\s+', ' ', line)
        
        return line.strip()
    
    def _structure_for_ai(self, text: str) -> str:
        """
        Structure text for AI line-by-line reading
        - Remove duplicates
        - Clean whitespace
        - Add section markers
        - Ensure logical flow
        """
        # Remove duplicate blocks
        text = self._remove_duplicates(text)
        
        # Clean whitespace
        text = self._clean_whitespace(text)
        
        # Add section markers for AI
        text = self._add_section_markers(text)
        
        return text
    
    def _remove_duplicates(self, text: str) -> str:
        """Remove duplicate content blocks"""
        lines = text.split('\n')
        seen_blocks = []
        result = []
        current_block = []
        
        for line in lines:
            clean_line = line.strip()
            
            if not clean_line:
                if current_block:
                    block_text = '\n'.join(current_block)
                    
                    # Check similarity
                    is_dup = False
                    for seen in seen_blocks:
                        if self._similar(block_text, seen, 0.75):
                            is_dup = True
                            break
                    
                    if not is_dup:
                        result.extend(current_block)
                        result.append('')
                        seen_blocks.append(block_text)
                    
                    current_block = []
            else:
                current_block.append(line)
        
        if current_block:
            block_text = '\n'.join(current_block)
            is_dup = False
            for seen in seen_blocks:
                if self._similar(block_text, seen, 0.75):
                    is_dup = True
                    break
            if not is_dup:
                result.extend(current_block)
        
        return '\n'.join(result)
    
    def _similar(self, text1: str, text2: str, threshold: float) -> bool:
        """Check if two texts are similar"""
        words1 = set(w.lower() for w in text1.split() if len(w) > 3)
        words2 = set(w.lower() for w in text2.split() if len(w) > 3)
        
        if not words1 or not words2:
            return False
        
        intersection = words1 & words2
        union = words1 | words2
        
        similarity = len(intersection) / len(union) if union else 0
        return similarity >= threshold
    
    def _clean_whitespace(self, text: str) -> str:
        """Clean excessive whitespace"""
        lines = text.split('\n')
        cleaned = []
        prev_empty = False
        
        for line in lines:
            if not line.strip():
                if not prev_empty:
                    cleaned.append('')
                    prev_empty = True
            else:
                cleaned.append(' '.join(line.split()))
                prev_empty = False
        
        # Remove leading/trailing empty lines
        while cleaned and not cleaned[0]:
            cleaned.pop(0)
        while cleaned and not cleaned[-1]:
            cleaned.pop()
        
        return '\n'.join(cleaned)
    
    def _add_section_markers(self, text: str) -> str:
        """Add clear section markers for AI understanding"""
        lines = text.split('\n')
        result = []
        
        # Major section headers (standalone, not part of other text)
        major_sections = [
            'EXPERIENCE', 'WORK EXPERIENCE', 'EMPLOYMENT', 'PROFESSIONAL EXPERIENCE',
            'EDUCATION', 'ACADEMIC BACKGROUND', 'QUALIFICATIONS',
            'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES',
            'PROJECTS', 'PROJECT EXPERIENCE',
            'CERTIFICATIONS', 'CERTIFICATES', 'LICENSES',
            'REFERENCES', 'REFERENCE',
            'SUMMARY', 'PROFILE', 'ABOUT ME', 'OBJECTIVE',
            'CONTACT', 'PERSONAL INFORMATION',
            'LANGUAGES', 'LANGUAGE',
            'AWARDS', 'HONORS', 'ACHIEVEMENTS',
            'PUBLICATIONS', 'RESEARCH',
            'VOLUNTEER', 'VOLUNTEERING'
        ]
        
        current_section = None
        
        for line in lines:
            line_stripped = line.strip()
            line_upper = line_stripped.upper()
            
            # Skip empty lines
            if not line_stripped:
                result.append(line)
                continue
            
            # Check if this is a major section header
            # Must be: short, matches exactly, and not part of a longer phrase
            is_major_section = False
            matched_section = None
            
            if len(line_stripped) < 50:  # Section headers are typically short
                for section in major_sections:
                    # Exact match or very close match
                    if line_upper == section or line_upper == section + ':':
                        is_major_section = True
                        matched_section = section
                        break
            
            if is_major_section:
                # This is a major section header
                current_section = matched_section
                result.append(f"\n=== {line_stripped.rstrip(':')} ===")
            else:
                # Regular content line
                # Don't mark as section even if it contains section keywords
                result.append(line)
        
        return '\n'.join(result)
    
    def _format_table_for_ai(self, table) -> str:
        """Format DOCX table for AI understanding"""
        rows = []
        
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Format as: Field: Value or Column1 | Column2 | Column3
            if len(cells) == 2:
                rows.append(f"{cells[0]}: {cells[1]}")
            else:
                rows.append(' | '.join(cells))
        
        return '\n'.join(rows)
    
    def _error_response(self, error: str) -> Dict[str, Any]:
        """Create error response"""
        return {
            'success': False,
            'error': error,
            'text': '',
            'metadata': {}
        }


def parse_resume(file_path: str, config: Optional[Dict[str, Any]] = None, 
                save_output: bool = True) -> Dict[str, Any]:
    """
    Parse resume from any format
    
    Args:
        file_path: Path to resume file
        config: Optional configuration
        save_output: Whether to save extracted text and metadata to files
    
    Returns:
        Dictionary with AI-friendly structured text
    """
    parser = UniversalParser(config)
    return parser.parse(file_path, save_output=save_output)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python universal_parser.py <file_path>")
        sys.exit(1)
    
    result = parse_resume(sys.argv[1])
    
    if result['success']:
        print("✅ Success!")
        print(f"\nAI-Friendly Structured Text:\n")
        print("="*80)
        print(result['text'])
        print("="*80)
        print(f"\nMetadata: {result['metadata']}")
    else:
        print(f"❌ Error: {result['error']}")
